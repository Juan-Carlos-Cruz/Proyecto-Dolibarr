#!/usr/bin/env python3
"""Generate a Word report aggregating Playwright (Allure) and PHPUnit results."""

from __future__ import annotations

import json
import os
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
ALLURE_RESULTS = ROOT / 'reports' / 'functional' / 'allure-results'
PHPUNIT_JUNIT = ROOT / 'reports' / 'phpunit' / 'phpunit-junit.xml'
OUTPUT = ROOT / 'docs' / 'informe-pruebas.docx'


def _load_allure_results() -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    if not ALLURE_RESULTS.exists():
        return results
    for file in ALLURE_RESULTS.glob('*-result.json'):
        with file.open(encoding='utf-8') as handler:
            try:
                results.append(json.load(handler))
            except json.JSONDecodeError:
                continue
    return results


def _summarize_allure(results: list[dict[str, Any]]) -> dict[str, Any]:
    status_counter = Counter(result.get('status', 'unknown') for result in results)
    duration = sum(result.get('time', {}).get('duration', 0) for result in results)
    return {
        'total': len(results),
        'status_counter': status_counter,
        'duration_ms': duration,
    }


def _load_phpunit_summary() -> dict[str, Any]:
    if not PHPUNIT_JUNIT.exists():
        return {'tests': 0, 'failures': 0, 'errors': 0}
    import xml.etree.ElementTree as ET

    tree = ET.parse(PHPUNIT_JUNIT)
    root = tree.getroot()
    return {
        'tests': int(root.attrib.get('tests', 0)),
        'failures': int(root.attrib.get('failures', 0)),
        'errors': int(root.attrib.get('errors', 0)),
    }


def build_document():
    document = Document()
    document.add_heading('Informe de Pruebas Automatizadas Dolibarr v22', 0)
    document.add_paragraph(f'Generado: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    allure_results = _load_allure_results()
    allure_summary = _summarize_allure(allure_results)
    document.add_heading('Resumen Selenium (Allure)', level=1)
    document.add_paragraph(f"Casos: {allure_summary['total']}")
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Estado'
    hdr_cells[1].text = 'Cantidad'
    for status, count in allure_summary['status_counter'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = str(count)
    document.add_paragraph(f"Duración total (ms): {allure_summary['duration_ms']}")

    phpunit_summary = _load_phpunit_summary()
    document.add_heading('Resumen PHPUnit', level=1)
    php_table = document.add_table(rows=1, cols=4)
    php_hdr = php_table.rows[0].cells
    php_hdr[0].text = 'Pruebas'
    php_hdr[1].text = 'Fallos'
    php_hdr[2].text = 'Errores'
    php_hdr[3].text = 'Estado'
    row = php_table.add_row().cells
    row[0].text = str(phpunit_summary['tests'])
    row[1].text = str(phpunit_summary['failures'])
    row[2].text = str(phpunit_summary['errors'])
    status = 'OK' if phpunit_summary['failures'] == 0 and phpunit_summary['errors'] == 0 else 'Revisar'
    row[3].text = status

    document.add_heading('Evidencias disponibles', level=1)
    document.add_paragraph('Allure raw: reports/functional/allure-results')
    document.add_paragraph('PHPUnit JUnit: reports/phpunit/phpunit-junit.xml')

    document.add_heading('Configuración de ejecución', level=1)
    env_table = document.add_table(rows=1, cols=2)
    env_table.rows[0].cells[0].text = 'Variable'
    env_table.rows[0].cells[1].text = 'Valor'
    for key in ['BASE_URL', 'ADMIN_USER', 'ADMIN_PASS', 'DOLIBARR_ROOT']:
        row = env_table.add_row().cells
        row[0].text = key
        row[1].text = os.getenv(key, 'N/D')

    document.add_heading('Notas', level=1)
    document.add_paragraph('Este reporte se genera automáticamente a partir de los resultados de Selenium (Pytest) y PHPUnit.')

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f'Reporte generado en {OUTPUT}')


if __name__ == '__main__':
    build_document()
